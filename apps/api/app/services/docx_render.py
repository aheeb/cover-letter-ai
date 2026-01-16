from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docxtpl import DocxTemplate, Listing
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

from app.models import LetterData


class TemplateNotFoundError(RuntimeError):
    pass


def render_letter_docx(
    *,
    template_path: Path,
    letter: LetterData,
    date_line: str,
    sender_adress: str = "",
    sender_name: str = "",
    location: str = "",
    recipient_indent_cm: float | None = None,
) -> bytes:
    if not template_path.exists():
        raise TemplateNotFoundError(f"Template not found: {template_path}")

    # Docxtpl supports paragraph breaks using the ASCII bell character (\a) via `Listing`.
    # This is the most reliable way to render multi-paragraph bodies into a single placeholder.
    body_listing = Listing("\a".join(letter.body_paragraphs))

    recipient_lines = [ln.strip() for ln in letter.recipient_block.splitlines() if ln.strip()]
    recipient_listing = Listing("\a".join(recipient_lines)) if recipient_lines else letter.recipient_block

    sender_lines = [ln.strip() for ln in sender_adress.splitlines() if ln.strip()]
    sender_listing = Listing("\a".join(sender_lines)) if sender_lines else sender_adress

    context: dict[str, object] = {
        # Template placeholders (as provided by user):
        "date": date_line,
        "recipient_address": recipient_listing,
        "role": letter.role_title,
        "body_of_motivational_letter": body_listing,
        "sender_adress": sender_listing,
        "sender_address": sender_listing,
        "sender_name": sender_name,
        "location": location,

        # Backwards-compatible keys (safe to keep):
        "date_line": date_line,
        "company": letter.company,
        "role_title": letter.role_title,
        "recipient_block": letter.recipient_block,
        "body_paragraphs": letter.body_paragraphs,
        "body_listing": body_listing,
    }

    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    rendered = BytesIO()
    tpl.save(rendered)
    rendered.seek(0)

    # Post-process: ensure recipient address block is consistently left-aligned, but positioned
    # on the right side of the page (common Swiss letter layout). This avoids template tweaks.
    doc = Document(rendered)

    # If we rely on the date line's tab stop for recipient alignment, capture it now
    # before we potentially move the date into a table.
    if recipient_indent_cm is None:
        paragraphs = list(_iter_all_paragraphs(doc))
        date_paragraph = _find_date_paragraph(paragraphs, date_line)
        tab_twips = _first_tab_stop_twips(date_paragraph) if date_paragraph is not None else None
        if tab_twips is not None:
            recipient_indent_cm = (tab_twips * 2.54) / 1440

    _format_sender_date_table(doc, sender_lines, location, date_line, recipient_indent_cm)
    _format_recipient_block(doc, recipient_lines, date_line, recipient_indent_cm)
    _format_sender_block(doc, sender_lines)

    final = BytesIO()
    doc.save(final)
    return final.getvalue()


def _format_recipient_block(
    doc: Document, recipient_lines: list[str], date_line: str, recipient_indent_cm: float | None
) -> None:
    if not recipient_lines:
        return

    paragraphs = list(_iter_all_paragraphs(doc))
    date_paragraph = _find_date_paragraph(paragraphs, date_line)
    indent = _recipient_block_indent(doc, date_paragraph, recipient_indent_cm)

    block = _find_paragraph_block(paragraphs, recipient_lines)
    if not _block_covers_all_recipient_lines(block, recipient_lines):
        # Fallback: locate the block by proximity/structure rather than exact text match.
        start_idx = 0
        if date_paragraph is not None:
            try:
                start_idx = paragraphs.index(date_paragraph) + 1
            except ValueError:
                start_idx = 0
        block = _find_recipient_block_by_proximity(paragraphs, recipient_lines, start_idx=start_idx)
    if not block:
        return

    # Prefer indent-based positioning when possible. This is more robust than tab stops and avoids cases
    # where only the first line inherits a leading tab (e.g. when Word keeps the block in one paragraph
    # with line breaks, or when additional paragraphs don't inherit the custom tab stops).
    if indent is not None:
        for p in block:
            _strip_leading_tab(p)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.left_indent = indent
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        return

    # Template has a leading TAB before `{{recipient_address}}` (verified in template.docx).
    # With docxtpl Listing, only the first line keeps that TAB → first line appears offset.
    # Fix: ensure *all* recipient lines start with the same TAB and avoid custom indents.
    if _starts_with_tab(block[0]):
        # Also ensure subsequent paragraphs keep the same *tab stop* definition as the first line.
        # Otherwise, extra lines can fall back to default tab stops and appear at the left margin.
        _copy_paragraph_tabs(block[0], block)
        for p in block:
            _ensure_leading_tab(p)
            p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.paragraph_format.left_indent = None
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
        return

    for p in block:
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = indent
        # Ensure first line has no extra indent (align all lines to the same left edge).
        p.paragraph_format.first_line_indent = Cm(0)
        # Tighten spacing to resemble Shift+Enter (line breaks) rather than paragraph gaps.
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def _format_sender_block(doc: Document, sender_lines: list[str]) -> None:
    if not sender_lines:
        return

    paragraphs = list(_iter_all_paragraphs(doc))
    block = _find_paragraph_block(paragraphs, sender_lines)
    if not _block_covers_all_recipient_lines(block, sender_lines):
        block = _find_recipient_block_by_proximity(paragraphs, sender_lines, start_idx=0)
    if not block:
        return

    for p in block:
        _strip_leading_tab(p)
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def _compose_date_line(location: str, date_line: str) -> str:
    loc = location.strip()
    date = date_line.strip()
    if loc and date:
        return f"{loc}, {date}"
    return loc or date


def _set_table_borders_none(table) -> None:
    tbl = table._tbl
    tbl_pr = getattr(tbl, "tblPr", None)
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    # Use fixed layout so column widths are respected consistently (Word + LibreOffice).
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = borders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            borders.append(edge_el)
        edge_el.set(qn("w:val"), "nil")


def _set_table_cell_margins_zero(table) -> None:
    """
    Remove default table cell margins/padding.

    Word defaults often add a small left/right inset inside cells. LibreOffice tends to preserve
    that inset when converting to PDF, which makes the right column text start slightly to the right
    of the intended column boundary. Setting margins to 0 makes the date block align exactly with
    the recipient block indent.
    """
    tbl = table._tbl
    tbl_pr = getattr(tbl, "tblPr", None)
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    mar = tbl_pr.find(qn("w:tblCellMar"))
    if mar is None:
        mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(mar)

    for edge in ("top", "left", "bottom", "right"):
        el = mar.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            mar.append(el)
        el.set(qn("w:w"), "0")
        el.set(qn("w:type"), "dxa")


def _format_sender_date_table(
    doc: Document,
    sender_lines: list[str],
    location: str,
    date_line: str,
    column_split_cm: float | None,
) -> None:
    """
    If sender address + date live in the same paragraph, replace it with a two-column
    table so the date aligns with the first sender line.
    """
    sender_lines = [ln.strip() for ln in sender_lines if ln and ln.strip()]
    if not sender_lines:
        return
    date_text = _compose_date_line(location, date_line)
    if not date_text:
        return

    paragraphs = list(_iter_all_paragraphs(doc))

    date_paragraph = _find_date_paragraph(paragraphs, date_line)
    if date_paragraph is None or _is_in_table_cell(date_paragraph):
        return

    try:
        date_idx = paragraphs.index(date_paragraph)
    except ValueError:
        return

    # docxtpl's Listing can expand `{{sender_adress}}` into multiple *paragraphs* while the
    # `{{location}}, {{date}}` run stays attached to the last one — so the date ends up aligned
    # with the LAST sender line. We detect the sender block preceding the date paragraph and
    # replace the whole cluster with a 2-column, borderless table.
    last_sender = sender_lines[-1]
    start_idx = None
    max_start = max(0, date_idx - (len(sender_lines) - 1))
    for candidate in range(max_start, date_idx + 1):
        ok = True
        for j, line in enumerate(sender_lines):
            pi = candidate + j
            if pi > date_idx:
                ok = False
                break
            p = paragraphs[pi]
            if pi == date_idx:
                # The date paragraph must contain both the date and the last sender line somewhere.
                if date_line not in p.text or line not in p.text:
                    ok = False
                    break
            else:
                if _normalize_recipient_text(p.text) != _normalize_recipient_text(line):
                    ok = False
                    break
        if ok:
            start_idx = candidate
            break

    if start_idx is None:
        # Fallback: if the date paragraph includes the last sender line, assume the sender block
        # starts at the first occurrence of the first sender line before the date.
        first_sender = sender_lines[0]
        for i in range(0, date_idx + 1):
            if first_sender in paragraphs[i].text:
                start_idx = i
                break
        if start_idx is None:
            return

    base_style = paragraphs[start_idx].style

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders_none(table)
    _set_table_cell_margins_zero(table)

    # Insert the table before the first paragraph of the cluster.
    paragraphs[start_idx]._p.addprevious(table._tbl)

    left_cell, right_cell = table.rows[0].cells
    left_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # Make the right column start exactly where the template's tab stop is (same position we
    # use for the recipient block indent), so the recipient block and date block align.
    if column_split_cm is not None:
        section = doc.sections[0]
        usable = section.page_width - (section.left_margin or Cm(2.0)) - (section.right_margin or Cm(2.0))
        left_w = Cm(column_split_cm)
        # Clamp: keep at least a little room for the right column.
        min_right = Cm(2.0)
        if left_w > usable - min_right:
            left_w = usable - min_right
        if left_w < Cm(2.0):
            left_w = Cm(2.0)
        right_w = usable - left_w
        table.columns[0].width = left_w
        table.columns[1].width = right_w

    # Fill left cell with sender lines (keeps exact line breaks).
    left_cell.text = sender_lines[0]
    lp = left_cell.paragraphs[0]
    lp.style = base_style
    lp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    lp.paragraph_format.left_indent = Cm(0)
    lp.paragraph_format.first_line_indent = Cm(0)
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)
    lp.paragraph_format.line_spacing = 1.0
    for line in sender_lines[1:]:
        p = left_cell.add_paragraph(line)
        p.style = base_style
        p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0

    # Fill right cell with date line, right-aligned.
    right_cell.text = date_text
    rp = right_cell.paragraphs[0]
    rp.style = base_style
    # Match the template behavior (date starts at the same tab stop as recipient block).
    rp.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    rp.paragraph_format.left_indent = Cm(0)
    rp.paragraph_format.first_line_indent = Cm(0)
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(0)
    rp.paragraph_format.line_spacing = 1.0

    # Remove the original sender/date cluster paragraphs to avoid duplicates.
    for pi in range(start_idx, date_idx + 1):
        p = paragraphs[pi]
        p._p.getparent().remove(p._p)


def _recipient_block_indent(doc: Document, date_paragraph, indent_override_cm: float | None) -> Emu | None:
    """
    Compute a default left indent that pushes the recipient block into the right page area.

    Override with `recipient_indent_cm` (float, centimeters) if you want to fine-tune.
    """
    if indent_override_cm is not None:
        return Cm(indent_override_cm)

    # If date line lives in a positioned container (e.g. table cell), don't add extra indent.
    if date_paragraph is not None and _is_in_table_cell(date_paragraph):
        return None

    # If the template uses a tab stop to position the date line, align the recipient block to the same
    # horizontal start so both blocks are perfectly aligned.
    if date_paragraph is not None:
        tab_twips = _first_tab_stop_twips(date_paragraph)
        if tab_twips is not None:
            # Word tab stops are in twips (1/1440 inch). 1 inch = 914400 EMU → 1 twip = 635 EMU.
            return Emu(int(tab_twips) * 635)

    # If the date line has an explicit indent, align recipient block to it.
    if date_paragraph is not None:
        li = date_paragraph.paragraph_format.left_indent
        if li is not None:
            return li  # type: ignore[return-value]

    section = doc.sections[0]
    left_margin = section.left_margin or Cm(2.0)
    right_margin = section.right_margin or Cm(2.0)
    usable = section.page_width - left_margin - right_margin

    # Default: start the block around ~60% into the usable width.
    return Emu(int(usable * 0.6))


def _iter_all_paragraphs(doc: Document):
    # Body paragraphs
    for p in doc.paragraphs:
        yield p
    # Tables in body
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    for row2 in nested.rows:
                        for cell2 in row2.cells:
                            for p2 in cell2.paragraphs:
                                yield p2


def _normalize_recipient_text(value: str) -> str:
    # Drop leading tabs (layout), normalize NBSP/ZWSP, collapse whitespace, and trim trailing commas.
    v = value.replace("\xa0", " ").replace("\u200b", " ").replace("\ufeff", " ")
    v = v.lstrip("\t")
    v = " ".join(v.split()).strip()
    v = v.rstrip(",").strip()
    return v


def _strip_leading_tab(paragraph) -> None:
    if not paragraph.runs:
        return
    t0 = paragraph.runs[0].text
    if t0.startswith("\t"):
        paragraph.runs[0].text = t0.lstrip("\t")


def _block_covers_all_recipient_lines(block, recipient_lines: list[str]) -> bool:
    if not block:
        return False
    targets = [_normalize_recipient_text(ln) for ln in recipient_lines if _normalize_recipient_text(ln)]
    if not targets:
        return False
    seen: set[str] = set()
    for p in block:
        for seg in p.text.split("\n"):
            n = _normalize_recipient_text(seg)
            if n:
                seen.add(n)
    return all(t in seen for t in targets)


def _find_recipient_block_by_proximity(paragraphs, recipient_lines: list[str], *, start_idx: int) -> list:
    """
    More robust block finder: locate the first paragraph containing the first recipient line, then
    include subsequent paragraphs until the next empty separator paragraph.
    """
    targets = [_normalize_recipient_text(ln) for ln in recipient_lines if _normalize_recipient_text(ln)]
    if not targets:
        return []
    first = targets[0]

    start = None
    for i in range(max(0, start_idx), len(paragraphs)):
        p = paragraphs[i]
        # Consider both whole-paragraph text and line-break-separated content.
        p_lines = [_normalize_recipient_text(x) for x in p.text.split("\n") if _normalize_recipient_text(x)]
        if not p_lines:
            continue
        if first in p_lines or first == _normalize_recipient_text(p.text):
            start = i
            break

    if start is None:
        return []

    block = []
    max_len = 12
    for j in range(start, min(len(paragraphs), start + max_len)):
        if j > start and _normalize_recipient_text(paragraphs[j].text) == "":
            break
        block.append(paragraphs[j])

    # Sanity: require we at least see the first line somewhere in the block.
    content_lines: set[str] = set()
    for p in block:
        for seg in p.text.split("\n"):
            n = _normalize_recipient_text(seg)
            if n:
                content_lines.add(n)
    if first not in content_lines:
        return []
    return block


def _find_paragraph_block(paragraphs, lines: list[str]):
    """
    Find the first consecutive paragraph block matching the given lines.

    Note: docxtpl + Word can introduce leading tabs, non-breaking spaces, and inconsistent whitespace
    between runs. We normalize aggressively so we don't miss lines and accidentally format only a
    subset of the recipient block (which can cause a line to jump to the left margin).
    """
    targets = [_normalize_recipient_text(ln) for ln in lines if _normalize_recipient_text(ln)]
    if not targets:
        return []

    texts = [_normalize_recipient_text(p.text) for p in paragraphs]
    n = len(targets)
    for i in range(0, max(0, len(texts) - n + 1)):
        if texts[i : i + n] == targets:
            return paragraphs[i : i + n]

    # Fallback: match all paragraphs that equal any of the target lines (keeps order)
    target_set = set(targets)
    return [p for p in paragraphs if _normalize_recipient_text(p.text) in target_set]


def _find_date_paragraph(paragraphs, date_line: str):
    needle = date_line.strip()
    if not needle:
        return None
    for p in paragraphs:
        text = p.text.strip()
        if text == needle or needle in text:
            return p
    return None


def _first_tab_stop_twips(paragraph) -> int | None:
    """
    Return the first <w:tab w:pos="..."> position (in twips) from the paragraph properties.
    """
    el = getattr(paragraph, "_p", None)
    if el is None:
        return None
    ppr = getattr(el, "pPr", None)
    if ppr is None:
        return None

    tabs_el = None
    for child in list(ppr):
        tag = getattr(child, "tag", "")
        if isinstance(tag, str) and tag.endswith("}tabs"):
            tabs_el = child
            break
    if tabs_el is None:
        return None

    for tab in list(tabs_el):
        tag = getattr(tab, "tag", "")
        if not (isinstance(tag, str) and tag.endswith("}tab")):
            continue
        # Attribute is usually namespaced: {...}pos
        for k, v in tab.attrib.items():
            if isinstance(k, str) and (k == "pos" or k.endswith("}pos")):
                try:
                    return int(v)
                except (TypeError, ValueError):
                    return None
    return None


def _is_in_table_cell(paragraph) -> bool:
    el = getattr(paragraph, "_p", None)
    if el is None:
        return False
    parent = el.getparent()
    while parent is not None:
        tag = getattr(parent, "tag", "")
        if isinstance(tag, str) and tag.endswith("}tc"):
            return True
        parent = parent.getparent()
    return False


def _starts_with_tab(paragraph) -> bool:
    text = paragraph.text
    if text.startswith("\t"):
        return True
    if paragraph.runs and paragraph.runs[0].text.startswith("\t"):
        return True
    return False


def _ensure_leading_tab(paragraph) -> None:
    """
    Ensure paragraph begins with a single tab character. This aligns to the same tab stop
    as the date line in the template without requiring Word edits.
    """
    if paragraph.runs:
        t0 = paragraph.runs[0].text
        if not t0.startswith("\t"):
            paragraph.runs[0].text = "\t" + t0.lstrip()
        else:
            paragraph.runs[0].text = "\t" + t0.lstrip("\t").lstrip()
        return
    paragraph.add_run("\t")


def _copy_paragraph_tabs(source_paragraph, paragraphs) -> None:
    """
    Copy the <w:tabs> paragraph-property element from the source paragraph onto all target paragraphs.

    This is needed because docxtpl's Listing may create new paragraphs for additional lines that do not
    inherit the original paragraph's custom tab stops, causing misalignment.
    """
    src_p = getattr(source_paragraph, "_p", None)
    if src_p is None:
        return
    src_ppr = getattr(src_p, "pPr", None)
    if src_ppr is None:
        return

    src_tabs = None
    for child in list(src_ppr):
        tag = getattr(child, "tag", "")
        if isinstance(tag, str) and tag.endswith("}tabs"):
            src_tabs = child
            break
    if src_tabs is None:
        return

    for p in paragraphs:
        tgt_p = getattr(p, "_p", None)
        if tgt_p is None:
            continue
        tgt_ppr = tgt_p.get_or_add_pPr()

        # Remove existing tabs (if any), then copy from source.
        for child in list(tgt_ppr):
            tag = getattr(child, "tag", "")
            if isinstance(tag, str) and tag.endswith("}tabs"):
                tgt_ppr.remove(child)
        tgt_ppr.append(deepcopy(src_tabs))


